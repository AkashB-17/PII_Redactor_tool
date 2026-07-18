"""
Automated PII Evaluator (corrected)
------------------------------------

...over the ENTIRE document's replacement mapping, checked against a
hand-picked SAMPLE ground truth (30 items on a 12,000-line document).
Every legitimately-detected PII instance that simply wasn't one of the
30 example items got counted as a false positive. That collapses
precision toward zero regardless of how good the detector actually is.

Fix: precision (and therefore FP) can only be measured honestly against
a ground truth that is EXHAUSTIVE for whatever text it's being compared
against. So this version scopes evaluation to a caller-specified prefix
of the document (the first `scope_blocks` paragraphs/table-cells, in
document order) -- and re-runs the detector directly on that scoped
text for the FP count, rather than trusting the whole-document
anonymizer mapping. You are responsible for making sure GROUND_TRUTH
in your test actually lists every PII instance within that scope, not
just interesting examples -- see the note in test_evaluator.py.

Secondary bug fixed: ground-truth entries that didn't literally
string-match the document (due to case/whitespace differences) were
silently dropped from both tp and fn, quietly shrinking the recall
denominator. This version normalizes whitespace/case before matching,
and explicitly reports any ground-truth entries it could not find in
the source document at all, instead of silently ignoring them.
"""

from __future__ import annotations

import re
from typing import List, Optional

from docx import Document


class AutomatedEvaluator:

    # ------------------------------------------------------------------
    # Text extraction
    # ------------------------------------------------------------------

    @staticmethod
    def extract_blocks_from_docx(file_path: str) -> List[str]:
        """
        Returns paragraph and table-cell text as a list of blocks, in
        document order (all paragraphs first, then all table cells).
        Kept as a list -- rather than one joined string -- so callers can
        slice a prefix ("first N blocks") consistently between an input
        and output document that share the same paragraph/cell structure.
        """
        doc = Document(file_path)
        blocks = []

        for p in doc.paragraphs:
            if p.text.strip():
                blocks.append(p.text.strip())

        for table in doc.tables:
            for row in table.rows:
                for cell in row.cells:
                    for p in cell.paragraphs:
                        if p.text.strip():
                            blocks.append(p.text.strip())

        return blocks

    @staticmethod
    def preview_blocks(file_path: str, n: int = 60) -> None:
        """
        Utility to help you pick a sensible `scope_blocks` value: prints
        the first n text blocks with their index, so you can see exactly
        where to stop labeling ground truth and set scope_blocks to match.
        """
        blocks = AutomatedEvaluator.extract_blocks_from_docx(file_path)
        for i, block in enumerate(blocks[:n]):
            print(f"[{i}] {block[:120]}")

    @staticmethod
    def normalize(text: str) -> str:
        return re.sub(r"\s+", " ", text.strip().lower())

    # ------------------------------------------------------------------
    # Evaluation
    # ------------------------------------------------------------------

    @staticmethod
    def evaluate(
        input_path: str,
        output_path: str,
        detector,
        ground_truth: List[str],
        scope_blocks: Optional[int] = None,
    ) -> dict:
        """
        detector: a PIIDetector instance (reuse the one already loaded by
                  your DocxProcessor -- pass processor.detector -- so we
                  don't load a second spaCy model).
        ground_truth: list of PII strings, EXHAUSTIVE for the scope being
                  evaluated (see scope_blocks).
        scope_blocks: restrict evaluation to the first N text blocks
                  (paragraphs + table cells, in document order). Must
                  match how far you actually went when building
                  ground_truth. If None, evaluates the whole document
                  and prints a loud warning, since an exhaustive
                  whole-document ground truth is unrealistic to hand-label.
        """
        input_blocks = AutomatedEvaluator.extract_blocks_from_docx(input_path)
        output_blocks = AutomatedEvaluator.extract_blocks_from_docx(output_path)

        if scope_blocks is not None:
            input_blocks = input_blocks[:scope_blocks]
            output_blocks = output_blocks[:scope_blocks]
            scope_label = f"first {scope_blocks} text blocks"
        else:
            scope_label = "entire document"
            print(
                "WARNING: no scope_blocks given -- evaluating over the whole "
                "document. Precision/False-Positive numbers are only "
                "meaningful if ground_truth genuinely lists EVERY PII "
                "instance in the whole document, not a sample."
            )

        original_text = " ".join(input_blocks)
        redacted_text = " ".join(output_blocks)

        norm_original = AutomatedEvaluator.normalize(original_text)
        norm_redacted = AutomatedEvaluator.normalize(redacted_text)
        norm_gold = {AutomatedEvaluator.normalize(g) for g in ground_truth}

        # ---- True Positives / False Negatives: did each gold PII item
        # actually disappear from the redacted output? ----
        tp, fn = 0, 0
        unmatched_ground_truth = []

        for entity in ground_truth:
            norm_entity = AutomatedEvaluator.normalize(entity)
            if norm_entity not in norm_original:
                # Couldn't even find this labeled entity in the source
                # document -- almost always a typo/spacing mismatch in
                # the ground truth itself. Report it instead of silently
                # dropping it from the denominator.
                unmatched_ground_truth.append(entity)
                continue
            if norm_entity not in norm_redacted:
                tp += 1
            else:
                fn += 1

        # ---- False Positives: re-run the detector directly on the SAME
        # scoped text used for ground truth, and check each detection
        # against the (exhaustive, for this scope) gold set. This is the
        # actual fix -- it no longer trusts the whole-document mapping. ----
        detected_entities = detector.detect(original_text)
        fp = 0
        false_positive_examples = []

        for ent in detected_entities:
            norm_value = AutomatedEvaluator.normalize(ent.value)
            # Substring match both ways to tolerate minor span-boundary
            # noise (trailing punctuation, extra whitespace) between what
            # the detector grabbed and how you transcribed the gold item.
            is_match = any(
                norm_value in gold or gold in norm_value
                for gold in norm_gold
            )
            if not is_match:
                fp += 1
                false_positive_examples.append((ent.entity_type, ent.value))

        precision = tp / (tp + fp) if (tp + fp) > 0 else 0.0
        recall = tp / (tp + fn) if (tp + fn) > 0 else 0.0
        f1_score = (
            2 * (precision * recall) / (precision + recall)
            if (precision + recall) > 0 else 0.0
        )

        if unmatched_ground_truth:
            print(
                f"WARNING: {len(unmatched_ground_truth)} ground-truth "
                f"entries were not found in the original document text "
                f"(check spelling/spacing) and were EXCLUDED from scoring: "
                f"{unmatched_ground_truth}"
            )

        return {
            "Scope": scope_label,
            "True Positives": tp,
            "False Positives": fp,
            "False Negatives": fn,
            "Unmatched Ground Truth (excluded)": len(unmatched_ground_truth),
            "Precision": round(precision, 4),
            "Recall": round(recall, 4),
            "F1-Score": round(f1_score, 4),
            "False Positive Examples (up to 15)": false_positive_examples[:15],
        }