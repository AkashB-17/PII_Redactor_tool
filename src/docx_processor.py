"""
DOCX Processor

Reads a DOCX document, detects PII,
replaces it with fake values and
writes a new DOCX.
"""

from pathlib import Path
from docx import Document
from src.detector import PIIDetector
from src.anonymizer import PIIAnonymizer
from src.utils import DetectedEntity

class DocxProcessor:

    def __init__(self):
        self.detector = PIIDetector()
        self.anonymizer = PIIAnonymizer()

    # ----------------------------------------------------
    # Replace helper
    # ----------------------------------------------------
    @staticmethod
    def replace_text_by_index(text: str, entities: list[DetectedEntity], mapping: dict[str, str]) -> str:
        """
        Replace detected entities using right-to-left index slicing.
        This prevents index shifting and stops the "Scunthorpe problem"
        where fake names accidentally get re-replaced.
        """
        # Sort entities descending by start index
        sorted_entities = sorted(entities, key=lambda e: e.start, reverse=True)
        
        for entity in sorted_entities:
            fake_value = mapping.get(entity.value, "<REDACTED>")
            # Replace exactly at the indices found by the detector
            text = text[:entity.start] + fake_value + text[entity.end:]
            
        return text

    def _process_paragraph(self, paragraph):
        """Helper to run detection and replacement on a single paragraph object."""
        # Note: Avoid using .strip() on extraction to retain document layout
        text = paragraph.text
        
        if not text.strip():
            return

        entities = self.detector.detect(text)

        if not entities:
            return

        mapping = self.anonymizer.build_mapping(entities)
        
        # Assign the rebuilt text back to the paragraph
        paragraph.text = self.replace_text_by_index(text, entities, mapping)

    # ----------------------------------------------------
    # Process document
    # ----------------------------------------------------
    def process(self, input_path: str, output_path: str):
        document = Document(input_path)

        # 1. Process standard paragraphs
        for paragraph in document.paragraphs:
            self._process_paragraph(paragraph)

        # 2. Process Tables (Crucial for Prospectus files)
        for table in document.tables:
            for row in table.rows:
                for cell in row.cells:
                    for paragraph in cell.paragraphs:
                        self._process_paragraph(paragraph)

        document.save(output_path)
        print(f"Saved redacted document -> {output_path}")